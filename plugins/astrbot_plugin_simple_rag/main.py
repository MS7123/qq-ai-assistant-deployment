from __future__ import annotations

import asyncio
import csv
import json
import math
import re
import time
import urllib.parse
import urllib.request
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from astrbot.api import logger
from astrbot.api.event import AstrMessageEvent, filter
from astrbot.api.star import Context, Star
from astrbot.core.utils.astrbot_path import get_astrbot_data_path


PLUGIN_NAME = "astrbot_plugin_simple_rag"
KNOWLEDGE_FILE = "knowledge.json"
MAX_CHUNK_SIZE = 500
CHUNK_OVERLAP = 80
TOP_K = 8
FULL_CONTEXT_CHUNK_LIMIT = 20
MAX_FILE_BYTES = 20 * 1024 * 1024
SUPPORTED_FILE_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".pdf",
    ".docx",
    ".xlsx",
    ".xlsm",
    ".csv",
    ".tsv",
}


@dataclass(frozen=True)
class KnowledgeChunk:
    chunk_id: str
    text: str
    source: str
    created_at: int


class SimpleRagPlugin(Star):
    def __init__(self, context: Context):
        super().__init__(context)
        self.data_dir = Path(get_astrbot_data_path()) / "plugin_data" / PLUGIN_NAME
        self.knowledge_file = self.data_dir / KNOWLEDGE_FILE
        self.recent_files: dict[str, list[str]] = {}
        self.data_dir.mkdir(parents=True, exist_ok=True)

    @filter.event_message_type(filter.EventMessageType.ALL, priority=50)
    async def record_file_message(self, event: AstrMessageEvent):
        """记录最近收到的文件，便于用户随后用 /learnfile 入库。"""
        file_refs = collect_file_refs(event)
        if not file_refs:
            return

        origin = event.unified_msg_origin or "unknown"
        self.recent_files[origin] = file_refs

        if not is_plugin_command(event.message_str):
            files = "\n".join(f"- {item}" for item in file_refs[:5])
            yield event.plain_result(
                "已收到文件引用。发送 /learnfile 可导入最近文件。\n"
                f"{files}"
            )

    @filter.command("learn")
    async def learn(self, event: AstrMessageEvent):
        """写入一段知识到本地知识库。用法：/learn 知识内容"""
        content = extract_command_body(event.message_str, "learn")
        if not content:
            yield event.plain_result("用法：/learn 知识内容")
            return

        chunks = chunk_text(content)
        existing = load_knowledge(self.knowledge_file)
        now = int(time.time())
        source = build_source(event)

        start_index = len(existing) + 1
        for offset, chunk in enumerate(chunks):
            existing.append(
                KnowledgeChunk(
                    chunk_id=f"k{start_index + offset}",
                    text=chunk,
                    source=source,
                    created_at=now,
                )
            )

        save_knowledge(self.knowledge_file, existing)
        yield event.plain_result(f"已写入 {len(chunks)} 条知识片段，当前共 {len(existing)} 条。")

    @filter.command("learnfile")
    async def learnfile(self, event: AstrMessageEvent):
        """读取 txt、md、pdf、docx 文件并写入知识库。用法：/learnfile 文件路径"""
        file_refs = collect_file_refs(event)
        body = extract_command_body(event.message_str, "learnfile")
        if body:
            file_refs.insert(0, body)
        if not file_refs:
            file_refs = self.recent_files.get(event.unified_msg_origin or "unknown", [])

        if not file_refs:
            yield event.plain_result(
                "用法：/learnfile /AstrBot/data/temp/example.pdf\n"
                "也可以先发送文件，再发送 /learnfile。"
            )
            return

        existing = load_knowledge(self.knowledge_file)
        now = int(time.time())
        source = build_source(event)
        start_index = len(existing) + 1
        added_chunks = 0
        imported_files: list[str] = []
        errors: list[str] = []

        for file_ref in file_refs:
            try:
                file_path = await asyncio.to_thread(
                    prepare_file_ref,
                    file_ref,
                    self.data_dir / "uploads",
                )
                text = await asyncio.to_thread(extract_text_from_file, file_path)
                chunks = chunk_text(text)
                if not chunks:
                    errors.append(f"{file_path.name}: 未提取到文本")
                    continue

                for offset, chunk in enumerate(chunks):
                    existing.append(
                        KnowledgeChunk(
                            chunk_id=f"k{start_index + added_chunks + offset}",
                            text=chunk,
                            source=f"{source}#{file_path.name}",
                            created_at=now,
                        )
                    )

                added_chunks += len(chunks)
                imported_files.append(file_path.name)
            except Exception as exc:
                logger.exception(f"Failed to import file into knowledge base: {file_ref}")
                errors.append(f"{file_ref}: {exc}")

        if added_chunks:
            save_knowledge(self.knowledge_file, existing)

        lines = [
            f"文件入库完成：新增 {added_chunks} 条片段，当前共 {len(existing)} 条。"
        ]
        if imported_files:
            lines.append("已导入：" + "、".join(imported_files))
        if errors:
            lines.append("问题：")
            lines.extend(f"- {error}" for error in errors[:5])

        yield event.plain_result("\n".join(lines))

    @filter.command("filedebug")
    async def filedebug(self, event: AstrMessageEvent):
        """查看当前消息中的文件组件字段，用于适配不同平台文件消息。"""
        yield event.plain_result(build_file_debug_report(event))

    @filter.command("ask")
    async def ask(self, event: AstrMessageEvent):
        """基于本地知识库检索并调用大模型回答。用法：/ask 问题"""
        question = extract_command_body(event.message_str, "ask")
        if not question:
            yield event.plain_result("用法：/ask 问题")
            return

        knowledge = load_knowledge(self.knowledge_file)
        if not knowledge:
            yield event.plain_result("知识库还是空的。请先使用 /learn 写入资料。")
            return

        if is_list_knowledge_request(question):
            yield event.plain_result(format_knowledge_list(knowledge))
            return

        matches = select_context(question, knowledge)
        if not matches:
            yield event.plain_result("没有检索到相关资料。可以先用 /learn 补充知识。")
            return

        try:
            answer = await self.generate_answer(event, question, matches)
        except Exception as exc:
            logger.exception(f"RAG answer generation failed: {exc}")
            fallback_context = format_context(matches)
            answer = (
                "大模型调用失败，先返回检索到的相关资料：\n\n"
                f"{fallback_context}\n\n"
                f"错误：{exc}"
            )

        yield event.plain_result(answer)

    @filter.command("kbstats")
    async def kbstats(self, event: AstrMessageEvent):
        """查看当前知识库统计信息。"""
        knowledge = load_knowledge(self.knowledge_file)
        total_chars = sum(len(item.text) for item in knowledge)
        yield event.plain_result(f"知识库统计：{len(knowledge)} 条片段，约 {total_chars} 个字符。")

    @filter.command("kblist")
    async def kblist(self, event: AstrMessageEvent):
        """列出知识库片段摘要。用法：/kblist [数量]"""
        knowledge = load_knowledge(self.knowledge_file)
        if not knowledge:
            yield event.plain_result("知识库还是空的。")
            return

        body = extract_command_body(event.message_str, "kblist")
        limit = parse_positive_int(body, default=20, maximum=100)
        yield event.plain_result(format_knowledge_list(knowledge, limit=limit))

    @filter.command("kbsearch")
    async def kbsearch(self, event: AstrMessageEvent):
        """调试检索结果。用法：/kbsearch 问题"""
        question = extract_command_body(event.message_str, "kbsearch")
        if not question:
            yield event.plain_result("用法：/kbsearch 问题")
            return

        knowledge = load_knowledge(self.knowledge_file)
        matches = retrieve(question, knowledge, TOP_K)
        if not matches:
            yield event.plain_result("没有检索到匹配片段。")
            return

        yield event.plain_result(format_context(matches))

    @filter.command("kbshow")
    async def kbshow(self, event: AstrMessageEvent):
        """查看某个知识片段全文。用法：/kbshow k3"""
        chunk_id = extract_command_body(event.message_str, "kbshow")
        if not chunk_id:
            yield event.plain_result("用法：/kbshow k3")
            return

        knowledge = load_knowledge(self.knowledge_file)
        chunk = find_chunk(knowledge, chunk_id)
        if chunk is None:
            yield event.plain_result(f"没有找到片段：{chunk_id}")
            return

        yield event.plain_result(
            f"{chunk.chunk_id} | {chunk.source}\n"
            f"created_at={chunk.created_at}\n\n"
            f"{chunk.text}"
        )

    @filter.command("kbclear")
    async def kbclear(self, event: AstrMessageEvent):
        """清空当前简易知识库。用法：/kbclear confirm"""
        body = extract_command_body(event.message_str, "kbclear")
        if body != "confirm":
            yield event.plain_result("此操作会清空知识库。确认请发送：/kbclear confirm")
            return

        save_knowledge(self.knowledge_file, [])
        yield event.plain_result("知识库已清空。")

    @filter.command("kbdelete")
    async def kbdelete(self, event: AstrMessageEvent):
        """按片段 ID 删除知识。用法：/kbdelete k3 k4"""
        body = extract_command_body(event.message_str, "kbdelete")
        chunk_ids = [item.strip().lower() for item in body.split() if item.strip()]
        if not chunk_ids:
            yield event.plain_result("用法：/kbdelete k3 k4")
            return

        knowledge = load_knowledge(self.knowledge_file)
        if not knowledge:
            yield event.plain_result("知识库还是空的。")
            return

        targets = set(chunk_ids)
        kept: list[KnowledgeChunk] = []
        deleted: list[KnowledgeChunk] = []
        for chunk in knowledge:
            if chunk.chunk_id.lower() in targets:
                deleted.append(chunk)
            else:
                kept.append(chunk)

        if not deleted:
            yield event.plain_result("没有找到要删除的片段：" + "、".join(chunk_ids))
            return

        save_knowledge(self.knowledge_file, kept)
        deleted_ids = "、".join(chunk.chunk_id for chunk in deleted)
        yield event.plain_result(
            f"已删除 {len(deleted)} 条片段：{deleted_ids}\n"
            f"当前剩余 {len(kept)} 条。"
        )

    @filter.command("kbdelete_source")
    async def kbdelete_source(self, event: AstrMessageEvent):
        """按来源删除知识。用法：/kbdelete_source example.pdf"""
        keyword = extract_command_body(event.message_str, "kbdelete_source").strip()
        if not keyword:
            yield event.plain_result("用法：/kbdelete_source example.pdf")
            return

        knowledge = load_knowledge(self.knowledge_file)
        if not knowledge:
            yield event.plain_result("知识库还是空的。")
            return

        kept: list[KnowledgeChunk] = []
        deleted: list[KnowledgeChunk] = []
        normalized_keyword = keyword.lower()
        for chunk in knowledge:
            if normalized_keyword in chunk.source.lower():
                deleted.append(chunk)
            else:
                kept.append(chunk)

        if not deleted:
            yield event.plain_result(f"没有找到来源包含 `{keyword}` 的片段。")
            return

        save_knowledge(self.knowledge_file, kept)
        yield event.plain_result(
            f"已删除来源包含 `{keyword}` 的 {len(deleted)} 条片段。\n"
            f"当前剩余 {len(kept)} 条。"
        )

    async def generate_answer(
        self,
        event: AstrMessageEvent,
        question: str,
        matches: list[tuple[KnowledgeChunk, float]],
    ) -> str:
        provider_id = await self.context.get_current_chat_provider_id(
            umo=event.unified_msg_origin
        )
        context_text = format_context(matches)
        prompt = (
            "请基于下面的知识库片段回答用户问题。"
            "如果片段中没有答案，请明确说明知识库资料不足，不要编造。\n\n"
            f"知识库片段：\n{context_text}\n\n"
            f"用户问题：{question}\n\n"
            "请用中文回答，并在末尾列出引用的片段编号。"
        )

        llm_resp = await self.context.llm_generate(
            chat_provider_id=provider_id,
            prompt=prompt,
        )
        completion_text = getattr(llm_resp, "completion_text", None)
        return completion_text or str(llm_resp) or "大模型没有返回文本。"

    async def terminate(self):
        logger.info("SimpleRagPlugin terminated.")


def extract_command_body(message: str, command: str) -> str:
    text = (message or "").strip()
    patterns = (f"/{command}", command)
    for prefix in patterns:
        if text == prefix:
            return ""
        if text.startswith(prefix + " "):
            return text[len(prefix) :].strip()
    return text


def is_plugin_command(message: str) -> bool:
    text = (message or "").strip()
    commands = (
        "/learn",
        "/learnfile",
        "/ask",
        "/kbstats",
        "/kblist",
        "/kbsearch",
        "/kbshow",
        "/kbdelete",
        "/kbdelete_source",
        "/kbclear",
        "/filedebug",
    )
    return any(text.startswith(command) for command in commands)


def build_source(event: AstrMessageEvent) -> str:
    get_sender_name = getattr(event, "get_sender_name", None)
    sender = get_sender_name() if callable(get_sender_name) else "unknown"
    origin = event.unified_msg_origin or "unknown"
    return f"{sender}@{origin}"


def is_list_knowledge_request(question: str) -> bool:
    normalized = re.sub(r"\s+", "", question.lower())
    keywords = (
        "所有知识片段",
        "全部知识片段",
        "列出知识片段",
        "知识库列表",
        "listchunks",
        "listknowledge",
    )
    return any(keyword in normalized for keyword in keywords)


def parse_positive_int(value: str, default: int, maximum: int) -> int:
    try:
        parsed = int(value.strip())
    except (TypeError, ValueError):
        return default
    return max(1, min(parsed, maximum))


def format_knowledge_list(chunks: list[KnowledgeChunk], limit: int = 20) -> str:
    shown = chunks[:limit]
    lines = [f"知识库片段列表：共 {len(chunks)} 条，显示前 {len(shown)} 条。"]
    for item in shown:
        preview = item.text.replace("\n", " ").strip()
        if len(preview) > 120:
            preview = preview[:117] + "..."
        lines.append(f"- {item.chunk_id} | {item.source} | {preview}")

    if len(chunks) > limit:
        lines.append(f"还有 {len(chunks) - limit} 条未显示，可用 /kblist 100 查看更多。")
    return "\n".join(lines)


def find_chunk(chunks: list[KnowledgeChunk], chunk_id: str) -> KnowledgeChunk | None:
    normalized = chunk_id.strip().lower()
    for chunk in chunks:
        if chunk.chunk_id.lower() == normalized:
            return chunk
    return None


def collect_file_refs(event: AstrMessageEvent) -> list[str]:
    message_obj = getattr(event, "message_obj", None)
    message_chain = getattr(message_obj, "message", []) if message_obj else []
    file_refs: list[str] = []

    for component in message_chain:
        file_refs.extend(extract_file_refs_from_component(component))

    for attr_name in ("raw_message", "message", "message_str"):
        raw_value = getattr(message_obj, attr_name, None) if message_obj else None
        file_refs.extend(extract_file_refs_from_raw(raw_value))

    return unique_items(file_refs)


def extract_file_refs_from_component(component: Any) -> list[str]:
    refs: list[str] = []
    component_type = component.__class__.__name__.lower()
    if "file" not in component_type and "document" not in component_type:
        return refs

    for attr_name in ("path", "file_path", "local_path", "url", "file", "name"):
        value = getattr(component, attr_name, None)
        if isinstance(value, str) and value.strip():
            refs.append(value.strip())

    refs.extend(extract_file_refs_from_raw(getattr(component, "__dict__", None)))
    return refs


def extract_file_refs_from_raw(raw_value: Any) -> list[str]:
    refs: list[str] = []
    if raw_value is None:
        return refs

    if isinstance(raw_value, dict):
        raw_type = str(raw_value.get("type", "")).lower()
        data = raw_value.get("data", raw_value)
        if raw_type in {"file", "document"} or looks_like_file_payload(data):
            refs.extend(extract_file_refs_from_mapping(data))
        for value in raw_value.values():
            refs.extend(extract_file_refs_from_raw(value))
    elif isinstance(raw_value, (list, tuple)):
        for item in raw_value:
            refs.extend(extract_file_refs_from_raw(item))
    elif isinstance(raw_value, str):
        refs.extend(extract_file_refs_from_text(raw_value))

    return refs


def looks_like_file_payload(value: Any) -> bool:
    if not isinstance(value, dict):
        return False
    keys = {str(key).lower() for key in value}
    return bool(keys & {"file", "path", "file_path", "local_path", "url"})


def extract_file_refs_from_mapping(value: Any) -> list[str]:
    if not isinstance(value, dict):
        return []

    refs: list[str] = []
    for key in ("path", "file_path", "local_path", "url", "file", "name"):
        item = value.get(key)
        if isinstance(item, str) and item.strip():
            refs.append(item.strip())
    return refs


def extract_file_refs_from_text(text: str) -> list[str]:
    refs: list[str] = []
    url_matches = re.findall(r"https?://[^\s\]\)\"']+", text)
    refs.extend(url_matches)

    path_matches = re.findall(
        r"(?:/AstrBot/data|data|temp|/tmp)[^\s\]\)\"']+\.(?:txt|md|markdown|pdf|docx|xlsx|xlsm|csv|tsv)",
        text,
        flags=re.IGNORECASE,
    )
    refs.extend(path_matches)
    return refs


def unique_items(items: list[str]) -> list[str]:
    seen: set[str] = set()
    unique: list[str] = []
    for item in items:
        normalized = item.strip()
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        unique.append(normalized)
    return unique


def build_file_debug_report(event: AstrMessageEvent) -> str:
    message_obj = getattr(event, "message_obj", None)
    message_chain = getattr(message_obj, "message", []) if message_obj else []
    refs = collect_file_refs(event)
    lines = ["文件调试信息：", f"识别到的文件引用：{refs or '无'}"]

    for index, component in enumerate(message_chain, start=1):
        lines.append(f"[{index}] type={component.__class__.__name__}")
        attrs = getattr(component, "__dict__", {})
        for key, value in attrs.items():
            value_text = repr(value)
            if len(value_text) > 300:
                value_text = value_text[:297] + "..."
            lines.append(f"  {key}={value_text}")

    return "\n".join(lines[:60])


def prepare_file_ref(file_ref: str, upload_dir: Path) -> Path:
    if is_url(file_ref):
        return download_file(file_ref, upload_dir)
    return resolve_local_file(file_ref)


def is_url(value: str) -> bool:
    parsed = urllib.parse.urlparse(value)
    return parsed.scheme in {"http", "https"}


def download_file(url: str, upload_dir: Path) -> Path:
    upload_dir.mkdir(parents=True, exist_ok=True)
    parsed = urllib.parse.urlparse(url)
    filename = Path(urllib.parse.unquote(Path(parsed.path).name)).name
    if not filename:
        filename = f"download-{int(time.time())}.txt"

    target = upload_dir / filename
    with urllib.request.urlopen(url, timeout=15) as response:
        content_length = response.headers.get("Content-Length")
        if content_length and int(content_length) > MAX_FILE_BYTES:
            raise ValueError(f"文件过大，当前限制 {MAX_FILE_BYTES // 1024 // 1024}MB")

        data = response.read(MAX_FILE_BYTES + 1)
        if len(data) > MAX_FILE_BYTES:
            raise ValueError(f"文件过大，当前限制 {MAX_FILE_BYTES // 1024 // 1024}MB")

    target.write_bytes(data)
    return target


def resolve_local_file(file_ref: str) -> Path:
    data_root = Path(get_astrbot_data_path()).resolve()
    path = Path(file_ref).expanduser()
    if not path.is_absolute():
        path = resolve_relative_file(path, data_root)

    resolved = path.resolve()
    if not is_relative_to(resolved, data_root):
        raise ValueError("只允许读取 /AstrBot/data 目录下的文件")
    if not resolved.exists() or not resolved.is_file():
        raise FileNotFoundError(f"文件不存在：{resolved}")
    if resolved.stat().st_size > MAX_FILE_BYTES:
        raise ValueError(f"文件过大，当前限制 {MAX_FILE_BYTES // 1024 // 1024}MB")
    return resolved


def resolve_relative_file(path: Path, data_root: Path) -> Path:
    candidates = [
        data_root / path,
        data_root / "temp" / path.name,
        data_root / "plugin_data" / PLUGIN_NAME / "uploads" / path.name,
    ]
    for candidate in candidates:
        if candidate.exists() and candidate.is_file():
            return candidate
    return candidates[0]


def is_relative_to(path: Path, parent: Path) -> bool:
    try:
        path.relative_to(parent)
        return True
    except ValueError:
        return False


def extract_text_from_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_FILE_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_FILE_EXTENSIONS))
        raise ValueError(f"暂不支持 {suffix or '无扩展名'} 文件，支持：{supported}")

    if suffix in {".txt", ".md", ".markdown"}:
        return read_plain_text(path)
    if suffix == ".pdf":
        return read_pdf_text(path)
    if suffix == ".docx":
        return read_docx_text(path)
    if suffix in {".xlsx", ".xlsm"}:
        return read_excel_text(path)
    if suffix in {".csv", ".tsv"}:
        return read_delimited_text(path, delimiter="\t" if suffix == ".tsv" else ",")

    raise ValueError(f"未实现的文件类型：{suffix}")


def read_plain_text(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="ignore")


def read_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        try:
            from PyPDF2 import PdfReader
        except ImportError as exc:
            raise RuntimeError("缺少 PDF 解析依赖，请安装 pypdf。") from exc

    reader = PdfReader(str(path))
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[page {index}]\n{text}")
    return "\n\n".join(pages)


def read_docx_text(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("缺少 Word 解析依赖，请安装 python-docx。") from exc

    document = Document(str(path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def read_excel_text(path: Path) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError("缺少 Excel 解析依赖，请安装 openpyxl。") from exc

    workbook = load_workbook(str(path), read_only=True, data_only=True)
    parts: list[str] = []
    try:
        for worksheet in workbook.worksheets:
            parts.append(f"[sheet {worksheet.title}]")
            for row in worksheet.iter_rows(values_only=True):
                values = [format_cell_value(value) for value in row]
                values = [value for value in values if value]
                if values:
                    parts.append(" | ".join(values))
    finally:
        workbook.close()

    return "\n".join(parts)


def read_delimited_text(path: Path, delimiter: str) -> str:
    text = read_plain_text(path)
    lines: list[str] = []
    reader = csv.reader(text.splitlines(), delimiter=delimiter)
    for row in reader:
        values = [cell.strip() for cell in row if cell.strip()]
        if values:
            lines.append(" | ".join(values))
    return "\n".join(lines)


def format_cell_value(value) -> str:
    if value is None:
        return ""
    if isinstance(value, float):
        return f"{value:g}"
    return str(value).strip()


def chunk_text(text: str) -> list[str]:
    normalized = re.sub(r"\s+", " ", text).strip()
    if not normalized:
        return []

    chunks: list[str] = []
    start = 0
    while start < len(normalized):
        end = min(start + MAX_CHUNK_SIZE, len(normalized))
        chunks.append(normalized[start:end].strip())
        if end == len(normalized):
            break
        start = max(end - CHUNK_OVERLAP, start + 1)
    return [chunk for chunk in chunks if chunk]


def load_knowledge(path: Path) -> list[KnowledgeChunk]:
    if not path.exists():
        return []

    try:
        raw_items = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        logger.exception(f"Failed to load knowledge file: {path}")
        return []

    chunks: list[KnowledgeChunk] = []
    for item in raw_items:
        try:
            chunks.append(
                KnowledgeChunk(
                    chunk_id=str(item["chunk_id"]),
                    text=str(item["text"]),
                    source=str(item.get("source", "unknown")),
                    created_at=int(item.get("created_at", 0)),
                )
            )
        except (KeyError, TypeError, ValueError):
            logger.warning(f"Skipped invalid knowledge item: {item}")
    return chunks


def save_knowledge(path: Path, chunks: list[KnowledgeChunk]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    payload = [
        {
            "chunk_id": item.chunk_id,
            "text": item.text,
            "source": item.source,
            "created_at": item.created_at,
        }
        for item in chunks
    ]
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def retrieve(
    question: str,
    chunks: list[KnowledgeChunk],
    top_k: int,
) -> list[tuple[KnowledgeChunk, float]]:
    query_terms = tokenize(question)
    if not query_terms:
        return []

    document_terms = [tokenize(chunk.text) for chunk in chunks]
    doc_count = len(chunks)
    document_frequency: dict[str, int] = {}
    for terms in document_terms:
        for term in set(terms):
            document_frequency[term] = document_frequency.get(term, 0) + 1

    scored: list[tuple[KnowledgeChunk, float]] = []
    for chunk, terms in zip(chunks, document_terms):
        score = bm25_like_score(query_terms, terms, document_frequency, doc_count)
        if score > 0:
            scored.append((chunk, score))

    scored.sort(key=lambda item: item[1], reverse=True)
    return scored[:top_k]


def select_context(
    question: str,
    chunks: list[KnowledgeChunk],
) -> list[tuple[KnowledgeChunk, float]]:
    if len(chunks) <= FULL_CONTEXT_CHUNK_LIMIT:
        matches = retrieve(question, chunks, len(chunks))
        scored_chunk_ids = {chunk.chunk_id for chunk, _ in matches}
        for chunk in chunks:
            if chunk.chunk_id not in scored_chunk_ids:
                matches.append((chunk, 0.0))
        return matches

    return retrieve(question, chunks, TOP_K)


def tokenize(text: str) -> list[str]:
    lowered = text.lower()
    terms = re.findall(r"[a-z0-9_]+|[\u4e00-\u9fff]", lowered)
    chinese_bigrams = [
        lowered[index : index + 2]
        for index in range(len(lowered) - 1)
        if is_chinese(lowered[index]) and is_chinese(lowered[index + 1])
    ]
    return terms + chinese_bigrams


def is_chinese(char: str) -> bool:
    return "\u4e00" <= char <= "\u9fff"


def bm25_like_score(
    query_terms: list[str],
    document_terms: list[str],
    document_frequency: dict[str, int],
    doc_count: int,
) -> float:
    if not document_terms:
        return 0.0

    term_counts: dict[str, int] = {}
    for term in document_terms:
        term_counts[term] = term_counts.get(term, 0) + 1

    score = 0.0
    doc_len_norm = 1 + math.log(1 + len(document_terms))
    for term in query_terms:
        tf = term_counts.get(term, 0)
        if tf == 0:
            continue
        df = document_frequency.get(term, 0)
        idf = math.log((doc_count + 1) / (df + 0.5)) + 1
        score += (tf / doc_len_norm) * idf
    return score


def format_context(matches: list[tuple[KnowledgeChunk, float]]) -> str:
    lines: list[str] = []
    for index, (chunk, score) in enumerate(matches, start=1):
        lines.append(
            f"[{index}] id={chunk.chunk_id}, score={score:.3f}, source={chunk.source}\n"
            f"{chunk.text}"
        )
    return "\n\n".join(lines)
